#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能文档工厂 - 批量证书生成器
作者：AI Assistant
功能：根据Excel名单批量生成Word证书
"""

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime

class CertificateGenerator:
    """证书生成器类"""
    
    def __init__(self, template_path=None):
        """
        初始化生成器
        :param template_path: Word模板路径，None则使用默认模板
        """
        self.template_path = template_path
        self.output_dir = "output"
        
        # 确保输出目录存在
        if not os.path.exists(self.output_dir):
            os.makedirs(self.output_dir)
    
    def create_default_template(self, template_name="template.docx"):
        """
        创建默认证书模板
        """
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1.5)
            section.bottom_margin = Inches(1.5)
            section.left_margin = Inches(1.5)
            section.right_margin = Inches(1.5)
        
        # 添加标题
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("荣誉证书")
        run.font.size = Pt(36)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
        
        # 添加空行
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 添加正文内容（带占位符）
        content = doc.add_paragraph()
        content.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content.add_run("兹证明 {name} 同志")
        run.font.size = Pt(16)
        
        doc.add_paragraph()
        
        content2 = doc.add_paragraph()
        content2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content2.add_run("在 {event} 中表现优异")
        run.font.size = Pt(16)
        
        doc.add_paragraph()
        
        content3 = doc.add_paragraph()
        content3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content3.add_run("特发此证，以资鼓励")
        run.font.size = Pt(16)
        
        # 添加空行
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()
        
        # 添加日期
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = date_para.add_run("颁发日期：{date}")
        run.font.size = Pt(14)
        
        # 添加颁发单位
        org_para = doc.add_paragraph()
        org_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = org_para.add_run("{organization}")
        run.font.size = Pt(14)
        run.font.bold = True
        
        # 保存模板
        template_path = os.path.join(self.output_dir, template_name)
        doc.save(template_path)
        print("[OK] 默认模板已创建: {}".format(template_path))
        return template_path
    
    def read_data(self, excel_path):
        """
        读取Excel数据
        :param excel_path: Excel文件路径
        :return: DataFrame
        """
        try:
            df = pd.read_excel(excel_path)
            print("[OK] 成功读取 {} 条数据".format(len(df)))
            print("[INFO] 数据列: {}".format(list(df.columns)))
            return df
        except Exception as e:
            print("[ERROR] 读取Excel失败: {}".format(e))
            return None
    
    def generate_certificates(self, excel_path, template_path=None, output_prefix="certificate"):
        """
        批量生成证书
        :param excel_path: Excel数据文件路径
        :param template_path: Word模板路径
        :param output_prefix: 输出文件名前缀
        """
        # 读取数据
        df = self.read_data(excel_path)
        if df is None:
            return False
        
        # 如果没有提供模板，创建默认模板
        if template_path is None or not os.path.exists(template_path):
            template_path = self.create_default_template()
        
        # 获取当前日期
        current_date = datetime.now().strftime("%Y年%m月%d日")
        
        # 批量生成
        generated_files = []
        for index, row in df.iterrows():
            try:
                # 读取模板
                doc = Document(template_path)
                
                # 替换占位符
                for paragraph in doc.paragraphs:
                    for run in paragraph.runs:
                        text = run.text
                        # 替换各种占位符
                        if '{name}' in text:
                            run.text = text.replace('{name}', str(row.get('姓名', row.get('name', '未知'))))
                        if '{event}' in text:
                            run.text = text.replace('{event}', str(row.get('活动', row.get('event', '本次活动'))))
                        if '{date}' in text:
                            run.text = text.replace('{date}', str(row.get('日期', row.get('date', current_date))))
                        if '{organization}' in text:
                            run.text = text.replace('{organization}', str(row.get('颁发单位', row.get('organization', '颁发单位'))))
                
                # 生成文件名
                name = str(row.get('姓名', row.get('name', '人员{}'.format(index+1))))
                output_filename = "{}_{}.docx".format(output_prefix, name)
                output_path = os.path.join(self.output_dir, output_filename)
                
                # 保存文档
                doc.save(output_path)
                generated_files.append(output_filename)
                print("[OK] 已生成: {}".format(output_filename))
                
            except Exception as e:
                print("[ERROR] 生成第 {} 个证书失败: {}".format(index+1, e))
        
        print("\n[DONE] 完成！共生成 {} 个证书".format(len(generated_files)))
        print("[PATH] 输出目录: {}".format(os.path.abspath(self.output_dir)))
        return True


def create_sample_excel():
    """
    创建示例Excel文件
    """
    data = {
        '姓名': ['张三', '李四', '王五', '赵六'],
        '活动': ['2026年度优秀员工评选', '2026年度优秀员工评选', '2026年度优秀员工评选', '2026年度优秀员工评选'],
        '颁发单位': ['XX科技有限公司', 'XX科技有限公司', 'XX科技有限公司', 'XX科技有限公司']
    }
    
    df = pd.DataFrame(data)
    output_path = "sample_data.xlsx"
    df.to_excel(output_path, index=False)
    print("[OK] 示例数据已创建: {}".format(output_path))
    return output_path


if __name__ == "__main__":
    print("=" * 50)
    print("智能文档工厂 - 证书生成器")
    print("=" * 50)
    
    # 创建示例数据
    print("\n[STEP 1] 创建示例数据...")
    excel_file = create_sample_excel()
    
    # 初始化生成器
    print("\n[STEP 2] 初始化生成器...")
    generator = CertificateGenerator()
    
    # 生成证书
    print("\n[STEP 3] 批量生成证书...")
    generator.generate_certificates(excel_file)
    
    print("\n" + "=" * 50)
    print("演示完成！")
    print("你可以修改 sample_data.xlsx 中的数据，然后重新运行此脚本")
    print("=" * 50)
