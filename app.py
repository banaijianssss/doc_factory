#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能文档工厂 - Web版
功能：网页端批量证书生成工具
"""

from flask import Flask, render_template, request, send_file, jsonify
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import io
import zipfile

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'

# 确保目录存在
for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)


def create_default_template():
    """创建默认证书模板"""
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.5)
        section.bottom_margin = Inches(1.5)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.5)
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("荣誉证书")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 正文
    content = doc.add_paragraph()
    content.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = content.add_run("兹证明 {name} 同志")
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    content2 = doc.add_paragraph()
    content2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = content2.add_run("在 {event} 中表现优异")
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    
    content3 = doc.add_paragraph()
    content3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = content3.add_run("特发此证，以资鼓励")
    run.font.size = Pt(16)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_para.add_run("颁发日期：{date}")
    run.font.size = Pt(14)
    
    # 颁发单位
    org_para = doc.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = org_para.add_run("{organization}")
    run.font.size = Pt(14)
    run.font.bold = True
    
    return doc


def generate_certificates_from_data(data_list, template_doc=None):
    """根据数据列表生成证书"""
    if template_doc is None:
        template_doc = create_default_template()
    
    current_date = datetime.now().strftime("%Y年%m月%d日")
    generated_files = []
    
    for index, row in enumerate(data_list):
        try:
            # 复制模板
            doc = Document()
            for element in template_doc.element.body:
                doc.element.body.append(element)
            
            # 替换占位符
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    text = run.text
                    if '{name}' in text:
                        run.text = text.replace('{name}', str(row.get('name', '未知')))
                    if '{event}' in text:
                        run.text = text.replace('{event}', str(row.get('event', '本次活动')))
                    if '{date}' in text:
                        run.text = text.replace('{date}', str(row.get('date', current_date)))
                    if '{organization}' in text:
                        run.text = text.replace('{organization}', str(row.get('organization', '颁发单位')))
            
            # 保存到内存
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            filename = f"certificate_{row.get('name', f'人员{index+1}')}.docx"
            generated_files.append((filename, output))
            
        except Exception as e:
            print(f"生成第 {index+1} 个证书失败: {e}")
    
    return generated_files


@app.route('/')
def index():
    """首页"""
    return render_template('index.html')


@app.route('/generate', methods=['POST'])
def generate():
    """生成证书API"""
    try:
        # 获取表单数据
        names = request.form.get('names', '').strip().split('\n')
        event = request.form.get('event', '本次活动')
        organization = request.form.get('organization', '颁发单位')
        date = request.form.get('date', datetime.now().strftime("%Y年%m月%d日"))
        
        # 清理空行
        names = [name.strip() for name in names if name.strip()]
        
        if not names:
            return jsonify({'error': '请输入至少一个姓名'}), 400
        
        # 构建数据列表
        data_list = []
        for name in names:
            data_list.append({
                'name': name,
                'event': event,
                'organization': organization,
                'date': date
            })
        
        # 生成证书
        generated_files = generate_certificates_from_data(data_list)
        
        if not generated_files:
            return jsonify({'error': '证书生成失败'}), 500
        
        # 打包成zip
        zip_output = io.BytesIO()
        with zipfile.ZipFile(zip_output, 'w', zipfile.ZIP_DEFLATED) as zf:
            for filename, file_data in generated_files:
                zf.writestr(filename, file_data.getvalue())
        
        zip_output.seek(0)
        
        return send_file(
            zip_output,
            mimetype='application/zip',
            as_attachment=True,
            download_name='certificates.zip'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/upload', methods=['POST'])
def upload():
    """上传Excel文件生成证书"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': '请上传文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择文件'}), 400
        
        # 读取Excel
        df = pd.read_excel(file)
        
        # 转换为字典列表
        data_list = df.to_dict('records')
        
        # 生成证书
        generated_files = generate_certificates_from_data(data_list)
        
        if not generated_files:
            return jsonify({'error': '证书生成失败'}), 500
        
        # 打包成zip
        zip_output = io.BytesIO()
        with zipfile.ZipFile(zip_output, 'w', zipfile.ZIP_DEFLATED) as zf:
            for filename, file_data in generated_files:
                zf.writestr(filename, file_data.getvalue())
        
        zip_output.seek(0)
        
        return send_file(
            zip_output,
            mimetype='application/zip',
            as_attachment=True,
            download_name='certificates.zip'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    print("=" * 50)
    print("智能文档工厂 - Web版")
    print("访问地址: http://localhost:5000")
    print("=" * 50)
    app.run(debug=True, host='0.0.0.0', port=5000)
