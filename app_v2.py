#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能文档工厂 - Web版 v2.0
功能：网页端批量证书生成工具（支持自定义模板）
"""

from flask import Flask, render_template, request, send_file, jsonify, session
import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
import io
import zipfile
import json

app = Flask(__name__)
app.secret_key = 'doc_factory_secret_key_2026'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024

UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'output'
TEMPLATE_FOLDER = 'custom_templates'

# 确保目录存在
for folder in [UPLOAD_FOLDER, OUTPUT_FOLDER, TEMPLATE_FOLDER]:
    if not os.path.exists(folder):
        os.makedirs(folder)

# 内置模板库
BUILTIN_TEMPLATES = {
    'honor': {
        'name': '荣誉证书',
        'description': '标准荣誉证书模板，适用于表彰、奖励场景',
        'preview': '🏆'
    },
    'completion': {
        'name': '结业证书',
        'description': '培训结业证书模板，适用于课程培训',
        'preview': '📜'
    },
    'participation': {
        'name': '参赛证明',
        'description': '活动参与证明模板，适用于比赛、活动',
        'preview': '🎖️'
    },
    'volunteer': {
        'name': '志愿者证书',
        'description': '志愿服务证明模板，适用于公益活动',
        'preview': '💝'
    },
    'employment': {
        'name': '聘书',
        'description': '聘任证书模板，适用于职位任命',
        'preview': '📋'
    },
    'scholarship': {
        'name': '奖学金证书',
        'description': '奖学金授予证书，适用于教育奖励',
        'preview': '🎓'
    }
}


def create_template(template_type='honor'):
    """创建不同类型的证书模板"""
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.2)
        section.bottom_margin = Inches(1.2)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.5)
    
    if template_type == 'honor':
        # 荣誉证书
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("荣誉证书")
        run.font.size = Pt(44)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xC4, 0x1E, 0x3A)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        content = doc.add_paragraph()
        content.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content.add_run("兹证明 {name} 同志")
        run.font.size = Pt(18)
        
        doc.add_paragraph()
        
        content2 = doc.add_paragraph()
        content2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content2.add_run("在 {event} 中表现优异")
        run.font.size = Pt(18)
        
        doc.add_paragraph()
        
        content3 = doc.add_paragraph()
        content3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content3.add_run("特发此证，以资鼓励")
        run.font.size = Pt(18)
        
    elif template_type == 'completion':
        # 结业证书
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("结业证书")
        run.font.size = Pt(44)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1A, 0x5F, 0x7A)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        content = doc.add_paragraph()
        content.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content.add_run("学员姓名：{name}")
        run.font.size = Pt(18)
        
        doc.add_paragraph()
        
        content2 = doc.add_paragraph()
        content2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content2.add_run("圆满完成 {event} 全部课程")
        run.font.size = Pt(18)
        
        doc.add_paragraph()
        
        content3 = doc.add_paragraph()
        content3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content3.add_run("成绩合格，准予结业")
        run.font.size = Pt(18)
        
    elif template_type == 'participation':
        # 参赛证明
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("参赛证明")
        run.font.size = Pt(44)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x2E, 0x7D, 0x32)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        content = doc.add_paragraph()
        content.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content.add_run("{name}")
        run.font.size = Pt(20)
        run.font.bold = True
        
        doc.add_paragraph()
        
        content2 = doc.add_paragraph()
        content2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content2.add_run("积极参加 {event}")
        run.font.size = Pt(18)
        
        doc.add_paragraph()
        
        content3 = doc.add_paragraph()
        content3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content3.add_run("特此证明")
        run.font.size = Pt(18)
        
    elif template_type == 'volunteer':
        # 志愿者证书
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("志愿者服务证书")
        run.font.size = Pt(40)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        content = doc.add_paragraph()
        content.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content.add_run("感谢 {name} 在")
        run.font.size = Pt(18)
        
        doc.add_paragraph()
        
        content2 = doc.add_paragraph()
        content2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content2.add_run("{event} 中")
        run.font.size = Pt(18)
        
        doc.add_paragraph()
        
        content3 = doc.add_paragraph()
        content3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content3.add_run("无私奉献，热心服务")
        run.font.size = Pt(18)
        
    elif template_type == 'employment':
        # 聘书
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("聘 书")
        run.font.size = Pt(48)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x8B, 0x45, 0x13)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        content = doc.add_paragraph()
        content.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content.add_run("兹聘请 {name} 为")
        run.font.size = Pt(18)
        
        doc.add_paragraph()
        
        content2 = doc.add_paragraph()
        content2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content2.add_run("{event}")
        run.font.size = Pt(20)
        run.font.bold = True
        
        doc.add_paragraph()
        
        content3 = doc.add_paragraph()
        content3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content3.add_run("聘期自 {date} 起生效")
        run.font.size = Pt(16)
        
    else:  # scholarship
        # 奖学金证书
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("奖学金证书")
        run.font.size = Pt(44)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0x8C, 0x00)
        
        doc.add_paragraph()
        doc.add_paragraph()
        
        content = doc.add_paragraph()
        content.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content.add_run("{name}")
        run.font.size = Pt(22)
        run.font.bold = True
        
        doc.add_paragraph()
        
        content2 = doc.add_paragraph()
        content2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content2.add_run("在 {event} 中成绩优异")
        run.font.size = Pt(18)
        
        doc.add_paragraph()
        
        content3 = doc.add_paragraph()
        content3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = content3.add_run("特授予奖学金，以资鼓励")
        run.font.size = Pt(18)
    
    # 通用底部
    for _ in range(3):
        doc.add_paragraph()
    
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = date_para.add_run("{date}")
    run.font.size = Pt(14)
    
    org_para = doc.add_paragraph()
    org_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = org_para.add_run("{organization}")
    run.font.size = Pt(14)
    run.font.bold = True
    
    return doc


def generate_certificates_from_data(data_list, template_type='honor', custom_template=None):
    """根据数据列表生成证书"""
    if custom_template:
        template_doc = Document(custom_template)
    else:
        template_doc = create_template(template_type)
    
    current_date = datetime.now().strftime("%Y年%m月%d日")
    generated_files = []
    
    for index, row in enumerate(data_list):
        try:
            doc = Document()
            for element in template_doc.element.body:
                doc.element.body.append(element)
            
            # 替换占位符
            for paragraph in doc.paragraphs:
                for run in paragraph.runs:
                    text = run.text
                    if '{name}' in text:
                        run.text = text.replace('{name}', str(row.get('name', row.get('姓名', '未知'))))
                    if '{event}' in text:
                        run.text = text.replace('{event}', str(row.get('event', row.get('活动', '本次活动'))))
                    if '{date}' in text:
                        run.text = text.replace('{date}', str(row.get('date', row.get('日期', current_date))))
                    if '{organization}' in text:
                        run.text = text.replace('{organization}', str(row.get('organization', row.get('颁发单位', '颁发单位'))))
            
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            filename = f"certificate_{row.get('name', row.get('姓名', f'人员{index+1}'))}.docx"
            generated_files.append((filename, output))
            
        except Exception as e:
            print(f"生成第 {index+1} 个证书失败: {e}")
    
    return generated_files


@app.route('/')
def index():
    """首页"""
    return render_template('index_v2.html', templates=BUILTIN_TEMPLATES)


@app.route('/api/templates')
def get_templates():
    """获取模板列表"""
    return jsonify(BUILTIN_TEMPLATES)


@app.route('/generate', methods=['POST'])
def generate():
    """生成证书API"""
    try:
        names = request.form.get('names', '').strip().split('\n')
        event = request.form.get('event', '本次活动')
        organization = request.form.get('organization', '颁发单位')
        date = request.form.get('date', datetime.now().strftime("%Y年%m月%d日"))
        template_type = request.form.get('template', 'honor')
        
        names = [name.strip() for name in names if name.strip()]
        
        if not names:
            return jsonify({'error': '请输入至少一个姓名'}), 400
        
        data_list = []
        for name in names:
            data_list.append({
                'name': name,
                'event': event,
                'organization': organization,
                'date': date
            })
        
        # 检查是否有自定义模板
        custom_template = None
        if 'template_file' in request.files:
            template_file = request.files['template_file']
            if template_file.filename:
                template_path = os.path.join(TEMPLATE_FOLDER, 'custom.docx')
                template_file.save(template_path)
                custom_template = template_path
        
        generated_files = generate_certificates_from_data(data_list, template_type, custom_template)
        
        if not generated_files:
            return jsonify({'error': '证书生成失败'}), 500
        
        zip_output = io.BytesIO()
        with zipfile.ZipFile(zip_output, 'w', zipfile.ZIP_DEFLATED) as zf:
            for filename, file_data in generated_files:
                zf.writestr(filename, file_data.getvalue())
        
        zip_output.seek(0)
        
        return send_file(
            zip_output,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'certificates_{template_type}.zip'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/upload', methods=['POST'])
def upload():
    """上传Excel文件生成证书"""
    try:
        if 'file' not in request.files:
            return jsonify({'error': '请上传文件'}), 400
        
        file = request.files['file']
        if file.filename == '':
            return jsonify({'error': '请选择文件'}), 400
        
        template_type = request.form.get('template', 'honor')
        
        df = pd.read_excel(file)
        data_list = df.to_dict('records')
        
        # 检查是否有自定义模板
        custom_template = None
        if 'template_file' in request.files:
            template_file = request.files['template_file']
            if template_file.filename:
                template_path = os.path.join(TEMPLATE_FOLDER, 'custom.docx')
                template_file.save(template_path)
                custom_template = template_path
        
        generated_files = generate_certificates_from_data(data_list, template_type, custom_template)
        
        if not generated_files:
            return jsonify({'error': '证书生成失败'}), 500
        
        zip_output = io.BytesIO()
        with zipfile.ZipFile(zip_output, 'w', zipfile.ZIP_DEFLATED) as zf:
            for filename, file_data in generated_files:
                zf.writestr(filename, file_data.getvalue())
        
        zip_output.seek(0)
        
        return send_file(
            zip_output,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f'certificates_{template_type}.zip'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/preview/<template_type>')
def preview_template(template_type):
    """预览模板"""
    try:
        doc = create_template(template_type)
        
        # 填充示例数据
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                text = run.text
                if '{name}' in text:
                    run.text = text.replace('{name}', '张三')
                if '{event}' in text:
                    run.text = text.replace('{event}', '2026年度优秀员工评选')
                if '{date}' in text:
                    run.text = text.replace('{date}', '2026年03月29日')
                if '{organization}' in text:
                    run.text = text.replace('{organization}', 'XX科技有限公司')
        
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=f'preview_{template_type}.docx'
        )
        
    except Exception as e:
        return jsonify({'error': str(e)}), 500


if __name__ == '__main__':
    print("=" * 50)
    print("智能文档工厂 - Web版 v2.0")
    print("访问地址: http://localhost:5000")
    print("=" * 50)
    app.run(debug=True, host='0.0.0.0', port=5000)
